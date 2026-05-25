from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Final

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS: Final[set[str]] = {".pdf", ".docx", ".txt"}
MAX_UPLOAD_BYTES: Final[int] = 8 * 1024 * 1024
MAX_EXTRACTED_CHARS: Final[int] = 20000


class DocumentExtractionError(ValueError):
    """Erro controlado para arquivos inválidos ou sem texto extraível."""


class DocumentExtractor:
    def extract(self, filename: str, content: bytes) -> str:
        if not filename:
            raise DocumentExtractionError("Arquivo sem nome.")

        if len(content) > MAX_UPLOAD_BYTES:
            raise DocumentExtractionError("Arquivo muito grande. Envie um arquivo de até 8 MB.")

        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise DocumentExtractionError(
                "Formato não suportado. Envie PDF, DOCX ou TXT. Arquivos .doc antigos não são suportados."
            )

        if extension == ".pdf":
            text = self._extract_pdf(content)
        elif extension == ".docx":
            text = self._extract_docx(content)
        else:
            text = self._extract_txt(content)

        normalized = self._normalize_text(text)
        if len(normalized) < 30:
            raise DocumentExtractionError(
                "Não foi possível extrair texto suficiente do arquivo. Verifique se o PDF não é apenas imagem/scan."
            )

        return normalized[:MAX_EXTRACTED_CHARS]

    def _extract_pdf(self, content: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except Exception as exc:  # noqa: BLE001 - resposta controlada para o usuário
            raise DocumentExtractionError("Não foi possível ler o PDF enviado.") from exc

    def _extract_docx(self, content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]

            table_cells: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_cells.append(cell.text)

            return "\n".join(paragraphs + table_cells)
        except Exception as exc:  # noqa: BLE001
            raise DocumentExtractionError("Não foi possível ler o DOCX enviado.") from exc

    def _extract_txt(self, content: bytes) -> str:
        for encoding in ("utf-8", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentExtractionError("Não foi possível decodificar o arquivo TXT.")

    def _normalize_text(self, text: str) -> str:
        lines = [" ".join(line.strip().split()) for line in text.splitlines()]
        return "\n".join(line for line in lines if line)
